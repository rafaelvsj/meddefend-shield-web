import mammoth
from docx import Document
import logging
from typing import Dict, Any
from .base_extractor import BaseExtractor

logger = logging.getLogger(__name__)

class DOCXExtractor(BaseExtractor):
    """Extract text from DOCX files using mammoth and python-docx."""
    
    async def extract(self, file_path: str, filename: str) -> Dict[str, Any]:
        """Extract text from DOCX using mammoth (preserves structure)."""
        
        try:
            # Try mammoth first (better formatting preservation)
            try:
                with open(file_path, "rb") as docx_file:
                    result = mammoth.extract_raw_text(docx_file)
                    text = result.value
                    
                if text and len(text.strip()) > 50:
                    logger.info(f"Extracted {len(text)} characters using mammoth")
                    return {
                        'text': text,
                        'method': 'mammoth',
                        'ocr_used': False
                    }
            except Exception as e:
                logger.warning(f"Mammoth extraction failed: {e}")
            
            # Fallback to python-docx
            doc = Document(file_path)
            paragraphs = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(" | ".join(row_data))
                
                if table_data:
                    paragraphs.append("\n".join(table_data))
            
            full_text = "\n\n".join(paragraphs)
            
            if len(full_text.strip()) < 10:
                raise Exception("No text content found in DOCX")
            
            logger.info(f"Extracted {len(full_text)} characters using python-docx")
            
            return {
                'text': full_text,
                'method': 'python-docx',
                'ocr_used': False
            }
            
        except Exception as e:
            logger.error(f"DOCX extraction failed: {str(e)}")
            raise Exception(f"DOCX extraction failed: {str(e)}")